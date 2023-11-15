from docx import Document

from docx.shared import RGBColor
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def write_peters_docx(info: dict):
    document = Document()

    document.add_paragraph(
        f"\t \t \t \t \t \t \t \t Prijs excl. BTW \t Netto prijs")

    heading = document.add_heading('', level=2)
    run = heading.add_run(info['title'])
    run.font.color.rgb = RGBColor(255, 0, 0)

    p = document.add_paragraph('')
    for s in info['specification']:
        p.add_run(f'{s}\n')

    # Creating a table object
    table = document.add_table(rows=0, cols=2)

    # looping through dict and adding data in rows
    for k, v in info['specs_table'].items():

        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(k)
        row[1].text = str(v)

    document.add_page_break()

    document.save(f'output/test.docx')
